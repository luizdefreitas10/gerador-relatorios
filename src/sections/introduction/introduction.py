from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import adicionar_titulo_secao, formatar_data_df 
from datetime import datetime
import pandas as pd

def gerar_secao_introducao(doc: Document, row):
    """
    Gera a seção '1. INTRODUÇÃO' do relatório, usando o utilitário de formatação de data.
    """
    adicionar_titulo_secao(doc, "1. INTRODUÇÃO")

    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    par.add_run(
        "A Coordenadoria de Transportes e Rodovias da Arpe realizou vistoria nos Terminais Rodoviários Intermunicipais concedidos com o objetivo de verificar as condições operacionais, de conservação, de manutenção e de segurança dos referidos terminais, conforme Contrato de Serviço Público "
    )
    
    par.add_run(str(row["Contrato"])).bold = True
    
    par.add_run(
        ", firmado entre o Governo do Estado, representado pela Secretaria de Transportes (SETRA) e a SOCICAM - Administração, Projetos e Representações Ltda. A ação foi no dia "
    )

    data_formatada = formatar_data_df(row["Data"]) 
    
    par.add_run(data_formatada).bold = True
    par.add_run(", exclusivamente no ")
    par.add_run(str(row["Local"])).bold = True

    periodo = str(row["Período"]).strip() if "Período" in row and pd.notna(row["Período"]) else ""
    if periodo:
       
        par.add_run(f" e nos dias {periodo}, nas cidades de xxxxxxxx, xxxxxxx, e xxxxxxx. ")
    else:
       
        par.add_run(" e nas cidades de xxxxxxxx, xxxxxxx, e xxxxxxx.. ")

    par.add_run("As visitas técnicas foram realizadas pela equipe formada por ")
    par.add_run(str(row["Pessoal Responsável"])).bold = True
    par.add_run(
        ".\n\nNeste Relatório de Fiscalização foram observadas as condições de conservação, limpeza e higiene das áreas de embarque e desembarque, dos sanitários, as condições do pavimento das vias de circulação interna, a infraestrutura oferecida, os locais de estocagem de veículos, a segurança e o atendimento ao usuário, bem como toda estrutura para funcionamento dos terminais. A equipe da Arpe conversou com os responsáveis pelos seis terminais que forneceram informações complementares à fiscalização, principalmente sobre a implantação dos sistemas contra incêndio."
    )