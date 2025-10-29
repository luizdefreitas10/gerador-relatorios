from docx import Document
from docx.shared import Pt, Inches
import pandas as pd
from tqdm import tqdm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os
import win32com.client as win32 

from sections.introduction.introduction import gerar_secao_introducao
from sections.objective.objective import gerar_secao_objetivo
from sections.legalbasis.legalbasis import gerar_secao_fundamentacao_legal
from sections.anexo.anexo import gerar_secao_anexo_fotos 
from sections.nonconformity.nonconformity import (
    gerar_secao_nao_conformidades_constatadas,
)
from sections.nonconformityresume.nonconformityresume import (
    gerar_secao_resumo_nao_conformidades,
)
from sections.finalconsiderations.finalconsiderations import (
    gerar_secao_consideracoes_finais,
)
from sections.summary.summary import inserir_quebra_e_sumario

from utils import (
    adicionar_texto_centralizado,
    adicionar_paragrafo_justificado,
    ajustar_largura_colunas,
    arquivo_em_uso,
    processar_imagem_para_relatorio,
    adicionar_imagem, 
)

def atualizar_toc_e_converter_para_pdf(caminho_docx, caminho_pdf):
    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(caminho_docx)
        
        doc.Fields.Update()
        doc.Save()
        doc.SaveAs2(caminho_pdf, FileFormat=17) 
        
        doc.Close(SaveChanges=False)
        word.Quit()
        return True
    
    except Exception as e:
        print(f"❌ ERRO ao gerar PDF/Atualizar Sumário (verifique se o Word está instalado): {e}")
        return False

def gerar_relatorio():
    """
    Gera o relatório completo (docx + pdf) com base nos dados da fiscalização.
    """

    if getattr(sys, "frozen", False):
        BASE_DIR = os.path.dirname(sys.executable)
    else:
        BASE_DIR = os.path.dirname(os.path.abspath(__file__))

    FOTOS_DIR = os.path.join(BASE_DIR, "assets")
    RELATORIOS_DIR = os.path.join(BASE_DIR, "reports")
    CAMINHO_PLANILHA = os.path.join(BASE_DIR, "planilha_fiscalizacao.xlsx")
    COLUNA_STATUS = "Relatório Gerado"

    os.makedirs(RELATORIOS_DIR, exist_ok=True)
    os.makedirs(FOTOS_DIR, exist_ok=True)

    if arquivo_em_uso(CAMINHO_PLANILHA):
        print("⚠️ A planilha está em uso. Feche-a antes de executar o script.")
        exit(1)
    
    print("\n--- Configuração de Pastas de Fotos ---")
    
    CAMINHO_RAIZ_FOTOS = None
    
    # Loop de validação do caminho das fotos (INÍCIO DA ALTERAÇÃO)
    while CAMINHO_RAIZ_FOTOS is None or not os.path.isdir(CAMINHO_RAIZ_FOTOS):
        
        # 1. Input para a pasta CTR-XX-YYYY (Ex: CTR-01-2025)
        while True:
            pasta_contrato = input("Digite a pasta CTR (ex: CTR-01-2025): ").strip()
            if pasta_contrato:
                break
            print("⚠️ O nome da pasta CTR é obrigatório.")

        # 2. Input para a pasta do Monitoramento (M#)
        while True:
            pasta_monitoramento = input("Digite a pasta do Monitoramento (ex: M1, M2): ").strip()
            if pasta_monitoramento:
                break
            print("⚠️ O nome da pasta de monitoramento é obrigatório.")
            
        # Constrói o caminho completo (Ex: assets/CTR-01-2025/M0)
        CAMINHO_RAIZ_FOTOS = os.path.join(FOTOS_DIR, pasta_contrato, pasta_monitoramento)
        
        if not os.path.isdir(CAMINHO_RAIZ_FOTOS):
            
            print(f"\n❌ ERRO DE CAMINHO: A pasta de evidências '{CAMINHO_RAIZ_FOTOS}' não foi encontrada.")
            print("Por favor, verifique se a pasta existe e se os inputs estão corretos, e digite novamente.")
            CAMINHO_RAIZ_FOTOS = None # Garante que o loop continue

    print(f"✅ Usando evidências de: {CAMINHO_RAIZ_FOTOS}")
    
    fiscalizacoes_df = pd.read_excel(CAMINHO_PLANILHA, sheet_name="Fiscalizações")
    nao_conformidades_df = pd.read_excel(
        CAMINHO_PLANILHA, sheet_name="Não-conformidades "
    )
    observacoes_df = pd.read_excel(
        CAMINHO_PLANILHA, sheet_name="Observações Importantes"
    )
    recomendacoes_df = pd.read_excel(
        CAMINHO_PLANILHA, sheet_name="Recomendações"
    )

    if COLUNA_STATUS not in fiscalizacoes_df.columns:
        fiscalizacoes_df[COLUNA_STATUS] = False
    fiscalizacoes_df[COLUNA_STATUS] = (
        fiscalizacoes_df[COLUNA_STATUS].fillna(False).astype(bool)
    )

    pendentes = fiscalizacoes_df[~fiscalizacoes_df[COLUNA_STATUS]]

    if pendentes.empty:
        print("✅ Nenhum relatório pendente.")
        return
        
    NOME_LOGO = "capa_monitoramento_arpe.jpg" 

    for idx in tqdm(pendentes.index, desc="Gerando relatórios"):
        row = fiscalizacoes_df.loc[idx]
        id_fisc = row["ID da Fiscalização"]
        doc = Document()
        
        section = doc.sections[0]
        section.top_margin = Inches(0.25)

        adicionar_texto_centralizado(doc, "COORDENADORIA DE TRANSPORTES E RODOVIAS")
    
        primeiro_paragrafo = doc.paragraphs[-1]
        primeiro_paragrafo.paragraph_format.space_before = Pt(0)
        primeiro_paragrafo.paragraph_format.space_after = Pt(0)

        doc.add_paragraph()

        adicionar_texto_centralizado(
            doc, "RELATÓRIO DO Xº MONITORAMENTO DO PROCESSO DE FICALIZAÇÃO TÉCNICO-OPERACIONAL CTR XX/XXXX"
        )

        doc.add_paragraph() 
        
        caminho_logo = os.path.join(FOTOS_DIR, NOME_LOGO)

        if os.path.exists(caminho_logo):
            buffer_logo = processar_imagem_para_relatorio(caminho_logo, largura_max=500, qualidade=95) 
            
            LARGURA_CAPA = Inches(6.5) 
            ALTURA_CAPA = Inches(5.5) 
            adicionar_imagem(
                doc, 
                buffer_logo, 
                largura_in=LARGURA_CAPA, 
                altura_in=ALTURA_CAPA
            )
            
        else:
            print(f"⚠️ Imagem de capa '{NOME_LOGO}' não encontrada. Pular logo.")
        
        adicionar_texto_centralizado(
            doc, "NÃO CONFORMIDADES NOS TERMINAIS RODOVIÁRIOS INTERMUNICIPAIS DE XXXXXXX, XXXXXXX E XXXXXXX CONCEDIDOS À EMPRESA SOCICAM"
        )
        adicionar_texto_centralizado(
            doc, "CONTRATO DE CONCESSÃO DE SERVIÇO PÚBLICO Nº X.XXX.XXX/XX"
        )
        adicionar_texto_centralizado(
            doc, "PROCESSO SEI Nº XXXXXXXXXX.XXXXXXXXX/XX"
        )
        
        doc.add_paragraph() 
        
        texto_data = "Recife, data de assinatura eletrônica"
        
        adicionar_paragrafo_justificado(doc, texto_data)
        
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        


        inserir_quebra_e_sumario(doc)

        gerar_secao_introducao(doc, row)
        gerar_secao_objetivo(doc)
        gerar_secao_fundamentacao_legal(doc)
        
        gerar_secao_nao_conformidades_constatadas(
            doc, row, nao_conformidades_df, FOTOS_DIR, observacoes_df, recomendacoes_df
        )
        
        gerar_secao_resumo_nao_conformidades(doc, row, nao_conformidades_df)
        
        gerar_secao_consideracoes_finais(doc, row)

        if nao_conformidades_df is not None:
            try:
               
                gerar_secao_anexo_fotos(doc, row, nao_conformidades_df, CAMINHO_RAIZ_FOTOS) 
            except Exception as e:
                
                print(f"Erro ao gerar Anexo: {e}")

      
        nome_arquivo = f"relatorio_{id_fisc}"
        caminho_docx = os.path.join(RELATORIOS_DIR, f"{nome_arquivo}.docx")
        caminho_pdf = os.path.join(RELATORIOS_DIR, f"{nome_arquivo}.pdf")

        doc.save(caminho_docx)
        
        sucesso_pdf = atualizar_toc_e_converter_para_pdf(caminho_docx, caminho_pdf)
       
        if os.path.exists(caminho_docx): 
            fiscalizacoes_df.at[idx, COLUNA_STATUS] = True

    if "Data" in fiscalizacoes_df.columns:
        fiscalizacoes_df["Data"] = pd.to_datetime(
            fiscalizacoes_df["Data"], errors="coerce"
        ).dt.strftime("%d/%m/%Y")

    if not arquivo_em_uso(CAMINHO_PLANILHA):
        with pd.ExcelWriter(
            CAMINHO_PLANILHA, engine="openpyxl", mode="a", if_sheet_exists="replace"
        ) as writer:
            fiscalizacoes_df.to_excel(writer, sheet_name="Fiscalizações", index=False)
            nao_conformidades_df.to_excel(
                writer, sheet_name="Não-conformidades ", index=False
            )
            observacoes_df.to_excel(
                writer, sheet_name="Observações Importantes", index=False
            )
            recomendacoes_df.to_excel(
                writer, sheet_name="Recomendações", index=False
            )

        ajustar_largura_colunas(CAMINHO_PLANILHA)
        
    print("🎉 Relatórios gerados e planilha atualizada com sucesso.")

    return

if __name__ == "__main__":
    gerar_relatorio()