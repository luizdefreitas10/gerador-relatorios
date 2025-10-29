from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING 
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT 
from openpyxl import load_workbook
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import io
from datetime import datetime
import pandas as pd
from docx.document import Document
from typing import Tuple

LARGURA_PADRAO_IN = Inches(6)
LARGURA_IMAGEM_LADO_A_LADO = Inches(3.25)
ALTURA_IMAGEM_LADO_A_LADO = Inches(2.7)

COR_CINZA_SOMBRA_HEX = "BFBFBF" 
COR_PRETO_RGB = (0, 0, 0)


def _set_run_language(run, lang_code="pt-BR"):
    """Define o idioma do texto (run) para evitar marcações falsas do corretor."""
    rPr = run._element.get_or_add_rPr()
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), lang_code)
    lang.set(qn('w:eastAsia'), lang_code)
    lang.set(qn('w:bidi'), lang_code)
    rPr.append(lang)

def desabilitar_correcao_paragrafo(paragrafo):
    """Desabilita a verificação ortográfica para um parágrafo inteiro.
    Usado em títulos e textos em caixa alta para eliminar falsos positivos."""
    pPr = paragrafo._element.get_or_add_pPr()
    
    no_proof = OxmlElement('w:noProof')
    pPr.append(no_proof)
    
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'pt-BR')
    pPr.append(lang)

def aplicar_sombreamento_paragrafo(paragrafo, cor_hex):
    """
    Aplica sombreamento (shading) ao parágrafo usando a cor hexadecimal.
    """
    pPr = paragrafo._element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear') 
    shading.set(qn('w:color'), 'auto') 
    shading.set(qn('w:fill'), cor_hex) 
    pPr.append(shading)


# FUNÇÕES DE ESTILO
def aplicar_estilo_texto(
    run, tamanho=12, negrito=False, fonte="Arial", cor_rgb=(0, 0, 0)
):
    """Aplica estilos básicos a um 'run' de texto."""
    run.font.name = fonte
    run._element.rPr.rFonts.set(qn("w:eastAsia"), fonte)
    run.font.size = Pt(tamanho)
    run.bold = negrito
    run.font.color.rgb = RGBColor(*cor_rgb)
    _set_run_language(run, "pt-BR")

def aplicar_estilo_corpo(run, negrito=False):
    """Aplica o estilo padrão do corpo do relatório (Arial 11pt)."""
    aplicar_estilo_texto(run, tamanho=11, negrito=negrito, fonte="Arial")
    

def aplicar_estilo_titulo(run, cor_rgb=COR_PRETO_RGB): 
    """Aplica o estilo padrão para títulos de seção (Arial 12pt, Negrito)."""
  
    aplicar_estilo_texto(run, tamanho=12, negrito=True, fonte="Arial", cor_rgb=cor_rgb)
    

# FUNÇÕES DE TEXTO/PARÁGRAFO (Novas e Corrigidas)
def adicionar_paragrafo_justificado(doc, texto, negrito=False):
    """Adiciona um parágrafo com texto justificado (Arial 11pt)."""
    paragrafo = doc.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    run = paragrafo.add_run(texto)
    aplicar_estilo_corpo(run, negrito=negrito) 
    return paragrafo


def adicionar_texto_centralizado(doc, texto, tamanho_fonte=12, negrito=True):
    """Adiciona um parágrafo com texto centralizado, com negrito (para a capa)."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(texto)
    
    if tamanho_fonte == 12 and negrito:
        aplicar_estilo_titulo(run)
    elif tamanho_fonte == 11:
        aplicar_estilo_corpo(run, negrito=negrito)
    else:
        aplicar_estilo_texto(run, tamanho_fonte, negrito, fonte="Arial")
        
    desabilitar_correcao_paragrafo(paragraph)
    
#FUNÇÕES DE COMPACTAÇÃO

def adicionar_paragrafo_info_compacta(doc: Document, titulo: str, conteudo: str):
    """Adiciona um parágrafo no formato Título: Conteúdo com espaçamento compactado."""
    par = doc.add_paragraph()
    
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
    par.paragraph_format.line_spacing = Pt(13) 
    
    run_titulo = par.add_run(titulo)
    aplicar_estilo_corpo(run_titulo, negrito=True) 
 
    run_conteudo = par.add_run(conteudo)
    aplicar_estilo_corpo(run_conteudo, negrito=False)
    
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    
    return par

def aplicar_estilo_paragrafo_compacto(paragrafo):
    """Aplica formatação compacta a qualquer parágrafo (usado em células de tabela)."""
    
    paragrafo.paragraph_format.space_before = Pt(0)
    paragrafo.paragraph_format.space_after = Pt(0)
    
    paragrafo.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
    paragrafo.paragraph_format.line_spacing = Pt(13)

# FUNÇÕES DE SEÇÃO (MODIFICADA)
def adicionar_titulo_secao(doc, texto, nivel_heading=1, cor_rgb=COR_PRETO_RGB, aplicar_sombra=False): 
    
    """
    Adiciona um título de seção formatado e usa doc.add_heading para o sumário.
    """

    try:
        par = doc.add_heading(level=nivel_heading)
    except Exception as e:
        print(f"Aviso: Erro ao aplicar Heading {nivel_heading}. Usando parágrafo padrão. Erro: {e}")
        par = doc.add_paragraph()

    run = par.add_run(texto)
    aplicar_estilo_titulo(run, cor_rgb=cor_rgb)

    if aplicar_sombra: 
        aplicar_sombreamento_paragrafo(par, COR_CINZA_SOMBRA_HEX)
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT 
        par.paragraph_format.space_after = Pt(6) 
    else:
        par.paragraph_format.space_after = Pt(6)
    
    run.font.all_caps = True
    
    desabilitar_correcao_paragrafo(par)


def adicionar_texto_contexto(doc: Document, contexto_tupla: Tuple[str, str]):
    """
    Adiciona o parágrafo de contexto formatado: TERMINAL (12pt, negrito, caixa alta) 
    seguido pelo restante do texto (11pt, normal).
    contexto_tupla deve ser (terminal, texto_restante)
    """
    terminal, texto_restante = contexto_tupla
    
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(12) 
    par.paragraph_format.space_after = Pt(6) 
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_terminal = par.add_run(terminal)
    aplicar_estilo_titulo(run_terminal)
    run_terminal.font.all_caps = True
    
    desabilitar_correcao_paragrafo(par)

    run_restante = par.add_run(texto_restante)
    aplicar_estilo_corpo(run_restante, negrito=False)

    return par


def adicionar_texto_esquerda(doc, texto, tamanho_fonte=11, negrito=False, compacto=False):
    """
    Adiciona um parágrafo com texto alinhado à esquerda.
    """
    paragraph = doc.add_paragraph()
    if compacto:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(texto)

    if tamanho_fonte == 11:
        aplicar_estilo_corpo(run, negrito=negrito)
    else:
        aplicar_estilo_texto(run, tamanho_fonte, negrito)

    return paragraph

def formatar_data_df(data_value):
    """
    Formata um valor de data para 'dd/mm/yyyy', lidando com tipos diferentes.
    """
    if pd.isna(data_value):
        return ""
    
    if isinstance(data_value, datetime):
        return data_value.strftime("%d/%m/%Y")
    
    if isinstance(data_value, str):
        try:
            return datetime.strptime(data_value.split()[0], "%Y-%m-%d").strftime("%d/%m/%Y")
        except ValueError:
            return data_value
            
    return str(data_value)


def ajustar_largura_colunas(caminho_planilha):
    wb = load_workbook(caminho_planilha)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for coluna in ws.columns:
            max_length = 0
            coluna_letra = coluna[0].column_letter

            for celula in coluna:
                try:
                    if celula.value:
                        max_length = max(max_length, len(str(celula.value)))
                except:
                    pass

            ajuste = max_length + 2
            ws.column_dimensions[coluna_letra].width = ajuste

    wb.save(caminho_planilha)

def arquivo_em_uso(caminho):
    try:
        os.rename(caminho, caminho)
        return False
    except PermissionError:
        return True

def aplicar_borda_paragrafo(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for border_name in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "2")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    pPr.append(borders)

def adicionar_legenda_formatada(doc, texto):
    par = doc.add_paragraph()
    run = par.add_run(texto)
    aplicar_estilo_texto(run, tamanho=10, fonte="Arial", cor_rgb=(90, 90, 90))
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aplicar_borda_paragrafo(par)

def processar_imagem_para_relatorio(caminho_imagem, largura_max=1024, qualidade=80):
    if not os.path.exists(caminho_imagem):
        print(f"ERRO DE FOTO (processar_imagem): Arquivo não encontrado: {caminho_imagem}")
        return None

    try:
        img = Image.open(caminho_imagem)
    except Exception as e:
        print(f"ERRO DE FOTO (processar_imagem): Não foi possível abrir '{caminho_imagem}'. Erro: {e}")
        return None

    if img.mode != "RGB":
        img = img.convert("RGB")

    if img.width > largura_max:
        proporcao = largura_max / float(img.width)
        altura_nova = int(float(img.height) * proporcao)
        img = img.resize((largura_max, altura_nova), Image.LANCZOS)

    buffer = io.BytesIO()
    img.save(buffer, format="JPEG", quality=qualidade, optimize=True)
    buffer.seek(0)
    return buffer

def adicionar_imagem(doc, buffer_imagem, largura_in=None, altura_in=None):
    """Adiciona uma imagem de um buffer ao documento. Centraliza o parágrafo."""
    largura_final = largura_in if largura_in is not None else LARGURA_PADRAO_IN
    paragrafo = doc.add_paragraph()

    if buffer_imagem:
        if altura_in is not None:
            paragrafo.add_run().add_picture(buffer_imagem, width=largura_final, height=altura_in)
        else:
            paragrafo.add_run().add_picture(buffer_imagem, width=largura_final)
    else:
        paragrafo.add_run("🚫 Imagem indisponível.")

    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def set_cell_border(cell, **kwargs):
    """Define as bordas da célula usando XML direto."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tblBorders = tcPr.first_child_found_in("w:tcBorders")
    if tblBorders is None:
        tblBorders = OxmlElement("w:tcBorders")
        tcPr.append(tblBorders)

    for border_name, attrs in kwargs.items():
        if border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border_element = OxmlElement(f'w:{border_name}')
            for k, v in attrs.items():
                border_element.set(qn(f'w:{k}'), str(v))

            old_border = tblBorders.find(qn(f'w:{border_name}'))
            if old_border is not None:
                tblBorders.remove(old_border)
                
            tblBorders.append(border_element)

def adicionar_legenda_formatada_na_celula(cell, texto):
    """Adiciona uma legenda formatada dentro de uma célula da tabela."""
    par = cell.paragraphs[0]

    if par.text.strip() != "":
        par.clear()

    run = par.add_run(texto)
    
    aplicar_estilo_texto(run, tamanho=10, fonte="Arial", cor_rgb=(90, 90, 90))
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT 

def adicionar_duas_imagens_lado_a_lado(doc: Document, fotos_dir: str, nome_foto1: str, legenda1: str, nome_foto2: str = None, legenda2: str = None, contexto_nc_tupla: Tuple[str, str, str] = None):
    """
    Adiciona duas imagens e suas legendas em uma tabela 2x1 (lado a lado).
    """
    if contexto_nc_tupla:
        terminal, nc_id, constatacao = contexto_nc_tupla
        
        texto_restante = f" - Não Conformidade {nc_id}: {constatacao}"
        adicionar_texto_contexto(doc, (terminal, texto_restante))

    tabela = doc.add_table(rows=2, cols=2)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.style = "Table Grid"

    col_width = LARGURA_IMAGEM_LADO_A_LADO

    for row_idx in range(2):
        for col_idx in range(2):
            cell = tabela.cell(row_idx, col_idx)
            cell.width = col_width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    caminho_imagem1 = os.path.join(fotos_dir, nome_foto1)
    buffer_img1 = processar_imagem_para_relatorio(caminho_imagem1)

    par1 = tabela.cell(0, 0).paragraphs[0]
    par1.alignment = WD_ALIGN_PARAGRAPH.CENTER 

    if buffer_img1 is not None:
        try:
            run1 = par1.add_run()
            run1.add_picture(buffer_img1, width=Inches(3.0), height=ALTURA_IMAGEM_LADO_A_LADO)
        except Exception as e:
            par1.add_run(f"Erro ao inserir foto 1: {e}.")
    else:
        par1.add_run(f"🚫 Imagem indisponível:{nome_foto1}")

    if nome_foto2:
        caminho_imagem2 = os.path.join(fotos_dir, nome_foto2)
        buffer_img2 = processar_imagem_para_relatorio(caminho_imagem2)

        par2 = tabela.cell(0, 1).paragraphs[0]
        par2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if buffer_img2 is not None:
            try:
                run2 = par2.add_run()
                run2.add_picture(buffer_img2, width=Inches(3.0), height=ALTURA_IMAGEM_LADO_A_LADO)
            except Exception as e:
               par2.add_run(f"⚠️ Erro ao inserir foto 2: {e}")
        else:
            par2.add_run(f"🚫 Imagem indisponível: {nome_foto2}")
    else:
       
        tabela.cell(0, 1).text = ""

    cell_legenda1 = tabela.cell(1, 0)
    cell_legenda1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
    adicionar_legenda_formatada_na_celula(cell_legenda1, legenda1)

    cell_legenda2 = tabela.cell(1, 1)
    cell_legenda2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if nome_foto2: 
        adicionar_legenda_formatada_na_celula(cell_legenda2, legenda2)
    else:
        cell_legenda2.text = ""
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12) 