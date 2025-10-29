from docx.document import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from utils import adicionar_texto_centralizado 
from docx.oxml import OxmlElement
from docx.oxml.ns import qn 

def inserir_quebra_e_sumario(doc: Document):
    """
    Insere uma quebra de página, adiciona o título 'SUMÁRIO'.
    """
   
    doc.add_section(WD_SECTION.NEW_PAGE)
 
    adicionar_texto_centralizado(doc, "SUMÁRIO", tamanho_fonte=12, negrito=True)
    
   
    doc.paragraphs[-1].paragraph_format.space_after = Pt(18)
    
    paragrafo_sumario = doc.add_paragraph()
    
    run_sumario_begin = paragrafo_sumario.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run_sumario_begin._element.append(fldChar_begin)
    
    # Instrução do campo: TOC \o "1-3" \h \z \u 
    run_sumario_instr = paragrafo_sumario.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    run_sumario_instr._element.append(instrText)

    run_sumario_end = paragrafo_sumario.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run_sumario_end._element.append(fldChar_end)

    doc.add_section(WD_SECTION.NEW_PAGE)