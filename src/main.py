import pandas as pd
from docx import Document
from docx.shared import Inches
import os
from docx2pdf import convert
from tqdm import tqdm
import sys
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from utils import (
    adicionar_paragrafo_justificado,
    adicionar_texto_centralizado,
    adicionar_titulo_secao,
    ajustar_largura_colunas,
    arquivo_em_uso,
)
from sections.introduction.introduction import gerar_secao_introducao
from sections.legalbasis.legalbasis import gerar_secao_fundamentacao_legal
from sections.nonconformity.nonconformity import (
    gerar_secao_nao_conformidades_constatadas,
)
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def main():
    if getattr(sys, "frozen", False):
        BASE_DIR = os.path.dirname(sys.executable)
    else:
        BASE_DIR = os.path.dirname(os.path.abspath(__file__))

    FOTOS_DIR = os.path.join(BASE_DIR, "assets")
    RELATORIOS_DIR = os.path.join(BASE_DIR, "reports")
    CAMINHO_PLANILHA = os.path.join(BASE_DIR, "planilha_fiscalizacao.xlsx")
    COLUNA_STATUS = "Relatório Gerado"

    os.makedirs(RELATORIOS_DIR, exist_ok=True)
    os.makedirs(FOTOS_DIR, exist_ok=True)

    if arquivo_em_uso(CAMINHO_PLANILHA):
        print("⚠️ A planilha está em uso. Feche-a antes de executar o script.")
        exit(1)

    fiscalizacoes_df = pd.read_excel(CAMINHO_PLANILHA, sheet_name="Fiscalizações")
    nao_conformidades_df = pd.read_excel(
        CAMINHO_PLANILHA, sheet_name="Não-conformidades "
    )

    if COLUNA_STATUS not in fiscalizacoes_df.columns:
        fiscalizacoes_df[COLUNA_STATUS] = False
    fiscalizacoes_df[COLUNA_STATUS] = (
        fiscalizacoes_df[COLUNA_STATUS].fillna(False).astype(bool)
    )

    pendentes = fiscalizacoes_df[~fiscalizacoes_df[COLUNA_STATUS]]

    if pendentes.empty:
        print("✅ Nenhum relatório pendente.")
        return

    for idx in tqdm(pendentes.index, desc="Gerando relatórios"):
        row = fiscalizacoes_df.loc[idx]
        id_fisc = row["ID da Fiscalização"]
        doc = Document()

        doc.add_picture(os.path.join(BASE_DIR, "assets/logo_arpe.png"), width=Inches(2))
        logo_arpe = doc.paragraphs[-1]
        logo_arpe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        adicionar_texto_centralizado(doc, "DIRETORIA DE REGULAÇÃO TÉCNICO-OPERACIONAL")
        adicionar_texto_centralizado(doc, "COORDENADORIA DE TRANSPORTES E RODOVIAS")
        adicionar_texto_centralizado(
            doc, "RELATÓRIO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL CTR 01/2025"
        )
        adicionar_texto_centralizado(
            doc, "TERMINAIS RODOVIÁRIOS INTERMUNICIPAIS CONCEDIDOS À EMPRESA SOCICAM"
        )
        adicionar_texto_centralizado(
            doc, "CONTRATO DE CONCESSÃO DE SERVIÇO PÚBLICO Nº 1.041.080/08"
        )

        doc.add_section(WD_SECTION.NEW_PAGE)

        gerar_secao_introducao(doc, row)

        gerar_secao_fundamentacao_legal(doc)

        gerar_secao_nao_conformidades_constatadas(
            doc, row, nao_conformidades_df, FOTOS_DIR
        )

        adicionar_titulo_secao(doc, "V - CONSTATAÇÕES")
        adicionar_paragrafo_justificado(
            doc, "A seguir, apresentam-se as não conformidades registradas:"
        )

        nc_fisc = nao_conformidades_df[
            nao_conformidades_df["ID da Fiscalização"] == id_fisc
        ]

        for nc_id, grupo_nc in nc_fisc.groupby("Nº"):
            descricao = grupo_nc["Não Conformidade"].iloc[0]
            doc.add_heading(f"{nc_id} - {descricao}", level=1)

            for _, linha in grupo_nc.iterrows():
                nome_foto = linha["Foto"]
                legenda = linha["Legenda da Foto"]
                foto_path = os.path.join(FOTOS_DIR, str(nome_foto))

                if os.path.exists(foto_path):
                    doc.add_picture(foto_path, width=Inches(3))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    adicionar_texto_centralizado(doc, legenda)
                else:
                    doc.add_paragraph(f"⚠️ Foto não encontrada: {nome_foto}")

        doc.add_section(WD_SECTION.NEW_PAGE)

        adicionar_titulo_secao(doc, "VII - CONCLUSÕES E RECOMENDAÇÕES")
        adicionar_paragrafo_justificado(
            doc, "Solicitamos plano de ação para regularização das não conformidades..."
        )
        adicionar_texto_centralizado(doc, f"\n\nRecife, {row['Data']}.")

        adicionar_texto_centralizado(doc, "\n\n_______________________________________")
        adicionar_texto_centralizado(doc, "Enildo Manoel da Silva Junior")
        adicionar_texto_centralizado(doc, "Analista de Regulação")

        caminho_docx = os.path.join(RELATORIOS_DIR, f"relatorio_{id_fisc}.docx")
        doc.save(caminho_docx)
        convert(caminho_docx, caminho_docx.replace(".docx", ".pdf"))
        fiscalizacoes_df.at[idx, COLUNA_STATUS] = True

    if not arquivo_em_uso(CAMINHO_PLANILHA):
        with pd.ExcelWriter(CAMINHO_PLANILHA, engine="openpyxl", mode="w") as writer:
            fiscalizacoes_df.to_excel(writer, sheet_name="Fiscalizações", index=False)
            nao_conformidades_df.to_excel(
                writer, sheet_name="Não-conformidades ", index=False
            )
        ajustar_largura_colunas(CAMINHO_PLANILHA)

    print("🎉 Relatórios gerados e planilha atualizada com sucesso.")


if __name__ == "__main__":
    try:
        main()
    finally:
        input("\nExecução finalizada. Pressione Enter para sair...")
