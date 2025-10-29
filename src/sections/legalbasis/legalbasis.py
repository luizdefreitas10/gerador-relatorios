from docx import Document
from docx.shared import Pt  
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import (
    adicionar_titulo_secao,
    aplicar_estilo_corpo,
    aplicar_estilo_paragrafo_compacto  
)

ESPACO_APOS_ITEM = Pt(12)


def _adicionar_item_legal(doc: Document, texto_formatado: list[tuple[str, bool]]):
    """
    Adiciona um item legal (lei, decreto, resolução) como um parágrafo justificado.
    """
    par = doc.add_paragraph()
    
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    
    aplicar_estilo_paragrafo_compacto(par)
    
    par.paragraph_format.space_after = ESPACO_APOS_ITEM
    
    run_hifen = par.add_run("- ")
    aplicar_estilo_corpo(run_hifen, negrito=False)

    for texto, negrito in texto_formatado:
        run = par.add_run(texto)
        aplicar_estilo_corpo(run, negrito=negrito)
    
    return par


def gerar_secao_fundamentacao_legal(doc: Document):
    """
    Gera a seção '3. FUNDAMENTAÇÃO LEGAL' no relatório, usando parágrafos separados
    para cada item para resolver o problema de justificação.
    """

    adicionar_titulo_secao(doc, "3. FUNDAMENTAÇÃO LEGAL")

    par_intro = doc.add_paragraph()
    par_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
   
    run_intro = par_intro.add_run(
        "A presente fiscalização encontra fundamento nas seguintes normas legais e regulamentares:"
    )
    aplicar_estilo_corpo(run_intro, negrito=False)
   
    par_intro.paragraph_format.space_after = Pt(12) 

    #ITEM 1: Lei nº 12.524, de 30 de dezembro de 2003 
    _adicionar_item_legal(
        doc,
        [
            ("Lei nº 12.524, de 30 de dezembro de 2003", True),
            (" – Altera e consolida as disposições da ", False),
            ("Lei nº 12.126, de 12 de dezembro de 2001", True),
            (
                ", que cria a Agência de Regulação dos Serviços Públicos Delegados do Estado de Pernambuco ARPE, regulamentada pelo ",
                False,
            ),
            ("Decreto nº 30.200, de 09 de fevereiro de 2007.", True),
        ],
    )

    #ITEM 2: Lei nº 13.254, de 21 de junho de 2007 
    _adicionar_item_legal(
        doc,
        [
            ("Lei nº 13.254, de 21 de junho de 2007", True),
            (" e alterações, em especial a ", False),
            ("Lei Estadual nº 15.200, de 17 de dezembro de 2013", True),
            (
                " Estrutura o Sistema de Transporte Coletivo Intermunicipal de Passageiros do Estado de Pernambuco, regulamentada pelo ",
                False,
            ),
            ("Decreto nº 40.559, de 31 de março de 2014.", True),
        ],
    )

    #ITEM 3: Resolução Arpe nº 46, de 07 de abril de 2008 
    _adicionar_item_legal(
        doc,
        [
            ("Resolução Arpe nº 46, de 07 de abril de 2008", True),
            (
                " (Antiga nº 06/2008) Aprova o Regulamento dos Terminais Rodoviários do Estado de Pernambuco, alterada parcialmente pela ",
                False,
            ),
            ("Resolução ARPE nº 53, de 26 de janeiro de 2009", True),
            (" (Antiga 003/2009).", False),
        ],
    )

    #ITEM 4: Resolução Arpe nº 083, de 30 de julho de 2013 
    _adicionar_item_legal(
        doc,
        [
            ("Resolução Arpe nº 083, de 30 de julho de 2013", True),
            (
                " Dispõe sobre os procedimentos de fiscalização, autuação e aplicação de penalidades aos prestadores de serviços públicos delegados no Estado de Pernambuco fiscalizados pela ARPE mediante delegação.",
                False,
            ),
        ],
    )

    #ITEM 5: Contrato de Concessão de Serviço Público nº 1.041.080/08 
    _adicionar_item_legal(
        doc,
        [
            ("Contrato de Concessão de Serviço Público nº 1.041.080/08, de 19 de setembro de 2008", True),
            (", e seus aditivos, especialmente o ", False),
            ("Segundo Termo Aditivo de 29 de setembro de 2017", True),
            (
                " contrato celebrado entre o Estado de Pernambuco, representado pela Secretaria de Transportes SETRA, e a SOCICAM Administração, Projetos e Representações Ltda.",
                False,
            ),
        ],
    )
