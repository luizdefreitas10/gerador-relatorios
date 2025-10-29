from utils import (
    adicionar_titulo_secao, 
    aplicar_estilo_corpo, 
    formatar_data_df, 
    aplicar_estilo_paragrafo_compacto, 
    adicionar_paragrafo_justificado
)
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.document import Document 

from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

def _aplicar_cor_fundo_celula(cell, cor_hex="D9D9D9"): 
    """Aplica sombreado (cor de fundo) à célula usando o valor HEX (Cinza Claro)."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), cor_hex)
    tcPr.append(shading)

def _aplicar_estilo_resumo(run, negrito=False):
    """Auxiliar para aplicar o estilo padrão da tabela de resumo (Arial 11pt)."""
    aplicar_estilo_corpo(run, negrito=negrito)
    
def _formatar_nome_terminal(nome_bruto):
    """Limpa e formata o nome do terminal."""
    if not isinstance(nome_bruto, str): return ""
    nome_maiusculo = nome_bruto.upper()
    nome_limpo = nome_maiusculo.replace("TERMINAL DE ", "").replace("TERMINAL DO ", "").replace("DE ", "").strip()
    return nome_limpo

def _extrair_sigla(nome_bruto):
    """Extrai a sigla ou as 3 primeiras letras do nome do terminal."""
    if not isinstance(nome_bruto, str): return ""

    if "(" in nome_bruto and ")" in nome_bruto:
        start = nome_bruto.find("(") + 1
        end = nome_bruto.find(")")
        if start < end:
            return nome_bruto[start:end].strip()
            
    return _formatar_nome_terminal(nome_bruto)[:3].upper() 


def gerar_secao_resumo_nao_conformidades(doc: Document, row, nao_conformidades_df):
    
    espaco1 = doc.add_paragraph()
    espaco1.paragraph_format.space_after = Pt(12)

    adicionar_titulo_secao(doc, "5. RESUMO DA SITUAÇÃO DAS NÃO CONFORMIDADES MONITORADAS")

    texto_introducao = (
        "O Quadro 1, a seguir resume os resultados das vistorias da equipe da Arpe nos Terminais Rodoviários "
        "Intermunicipais concedidos à SOCICAM, no período (DATA), nas cidades de "
        "(TERMINAIS). Foi constatado que as QUANTIDADE (NÚMERO DA QUANTIDADE) Não Conformidades "
        "apontadas no Relatório de Fiscalização Técnico-Operacional Arpe/CTR nº 02/2024 foram solucionadas pela "
        "Concessionária. "
    )
   
    adicionar_paragrafo_justificado(doc, texto_introducao)

    par_quadro = doc.add_paragraph()
    par_quadro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    par_quadro.paragraph_format.space_before = Pt(12) 
    par_quadro.paragraph_format.space_after = Pt(6)   
    
    run_parte1 = par_quadro.add_run("Quadro 1 - ")
    aplicar_estilo_corpo(run_parte1, negrito=True)

    run_parte2 = par_quadro.add_run("Resumo da Situação das Não Conformidades Pendentes")
    aplicar_estilo_corpo(run_parte2, negrito=True) 

    run_parte3 = par_quadro.add_run(" - RELATÓRIO ARPE/CTR 02/2024")
    aplicar_estilo_corpo(run_parte3, negrito=False)

    id_fisc = row["ID da Fiscalização"]

    nc_fisc = nao_conformidades_df[
        nao_conformidades_df["ID da Fiscalização"] == id_fisc
    ].copy() 

    CHAVE_INF_SOCICAM = "Informação SOCICAM carta"
    CHAVE_DATA_VISTORIA = "Vistoria da Arpe"
    CHAVE_SITUACAO = "Situação"

    if nc_fisc.empty:
        doc.add_paragraph("Nenhuma não conformidade registrada.")
        return


    tabela = doc.add_table(rows=1, cols=5) 
    tabela.style = "Table Grid" 
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER 
    
    col_widths = [Inches(1.0), Inches(2.2), Inches(1.9), Inches(1.5), Inches(0.8)] 

    cabecalho = tabela.rows[0].cells
    
    cabecalho[0].text = "TERMINAL"
    cabecalho[1].text = "NÃO CONFORMIDADE\nRELATÓRIO ARPE/CTR\nXX/XXXX" 
    cabecalho[2].text = "INFORMAÇÃO SOCICAM\nCarta SAP/PER/ARPE\nXXX/XXXX" 
    cabecalho[3].text = "VISTORIA DA ARPE\n[DATAS]" 
    cabecalho[4].text = "SITUAÇÃO"
    
    for i, cell in enumerate(cabecalho):
        _aplicar_cor_fundo_celula(cell, cor_hex="D9D9D9")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
        cell.width = col_widths[i]
        
        for par in cell.paragraphs:
            if par.runs:
                _aplicar_estilo_resumo(par.runs[0], negrito=True) 
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER

            
    #Preenchimento dos Dados
    current_word_row_index = 1 

    for terminal_bruto, grupo in nc_fisc.groupby("Terminal"):
        
        start_row_index = current_word_row_index 
        primeira_celula_terminal = None 
        
        grupo = grupo.sort_values(by="Nº")

        nome_terminal_limpo = _formatar_nome_terminal(terminal_bruto)
        
        for idx, (_, linha) in enumerate(grupo.iterrows()):
            
            row_cells = tabela.add_row().cells
            
            end_row_index = current_word_row_index 

            for i, cell in enumerate(row_cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
                cell.width = col_widths[i]
            
            if idx == 0:
                primeira_celula_terminal = row_cells[0]
                par = primeira_celula_terminal.paragraphs[0]
                
                run_terminal = par.add_run(nome_terminal_limpo)
                _aplicar_estilo_resumo(run_terminal, negrito=True) 
                
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                row_cells[0].text = ""

          
            numero_nc = str(linha.get("Nº", "")).strip() 
            desc_nc = str(linha.get("Não Conformidade", "")).strip()

            paragrafo_nc = row_cells[1].paragraphs[0]
     
            aplicar_estilo_paragrafo_compacto(paragrafo_nc)
            
            paragrafo_nc.alignment = WD_ALIGN_PARAGRAPH.LEFT 
      
            run_num_nc = paragrafo_nc.add_run(f"{_extrair_sigla(terminal_bruto)} {numero_nc}")
            _aplicar_estilo_resumo(run_num_nc, negrito=True) 
            
            run_desc_nc = paragrafo_nc.add_run(f" - {desc_nc}")
            _aplicar_estilo_resumo(run_desc_nc, negrito=False) 

            inf_socicam = str(linha.get(CHAVE_INF_SOCICAM, "N/A")).strip() 
            paragrafo_inf = row_cells[2].paragraphs[0]
            paragrafo_inf.alignment = WD_ALIGN_PARAGRAPH.LEFT 
            run_inf = paragrafo_inf.add_run(inf_socicam)
            _aplicar_estilo_resumo(run_inf, negrito=False)

            data_raw = linha.get(CHAVE_DATA_VISTORIA)
            data_vistoria = formatar_data_df(data_raw) if data_raw else "N/A"
            
            paragrafo_data = row_cells[3].paragraphs[0]
            paragrafo_data.alignment = WD_ALIGN_PARAGRAPH.LEFT 
            run_data = paragrafo_data.add_run(data_vistoria)
            _aplicar_estilo_resumo(run_data, negrito=False)

            situacao = str(linha.get(CHAVE_SITUACAO, "N/A")).upper().strip()
            
            negrito_sit = situacao in ('PENDENTE', 'NÃO CONFORME', 'NAO CONFORME') 

            paragrafo_sit = row_cells[4].paragraphs[0]
            paragrafo_sit.alignment = WD_ALIGN_PARAGRAPH.CENTER 
            run_sit = paragrafo_sit.add_run(situacao)
            _aplicar_estilo_resumo(run_sit, negrito=negrito_sit)
            
            current_word_row_index += 1 
      
        if grupo.shape[0] > 1 and primeira_celula_terminal:
            ultima_celula_terminal = tabela.cell(end_row_index, 0)
            primeira_celula_terminal.merge(ultima_celula_terminal)

    espaco = doc.add_paragraph()
    espaco.paragraph_format.space_after = Pt(24)