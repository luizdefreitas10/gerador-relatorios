from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import (
    adicionar_paragrafo_justificado,
    adicionar_titulo_secao,
    adicionar_texto_esquerda,
    aplicar_estilo_corpo, 
    aplicar_estilo_titulo, 
    adicionar_paragrafo_info_compacta 
)

def _inserir_texto_nc(doc, num_terminal, terminal, nc_titulo_identificador, descricao):
    """
    Insere o cabeçalho de texto de uma Não Conformidade.
    """
    
    par_terminal = doc.add_paragraph()
    run_terminal = par_terminal.add_run(f"4.{num_terminal} - {terminal.upper()}")
    
    aplicar_estilo_titulo(run_terminal)
    
    par_terminal.paragraph_format.space_before = Pt(12)
    par_terminal.paragraph_format.space_after = Pt(6)

    par_nc = doc.add_paragraph()
    
    run_nc_titulo = par_nc.add_run(f"Não Conformidade {nc_titulo_identificador}") 
    
    aplicar_estilo_corpo(run_nc_titulo, negrito=True)
    run_nc_titulo.underline = True
    run_nc_titulo.font.color.rgb = RGBColor(0, 0, 0) 
    
    run_nc_traco = par_nc.add_run(" - ")
   
    aplicar_estilo_corpo(run_nc_traco, negrito=False)
    run_nc_traco.underline = False
    run_nc_traco.font.color.rgb = RGBColor(0, 0, 0) 

    run_nc_desc = par_nc.add_run(descricao) 
    
    aplicar_estilo_corpo(run_nc_desc, negrito=False) 
    run_nc_desc.underline = False 
    run_nc_desc.font.color.rgb = RGBColor(0, 0, 0) 
    
    par_nc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW

def _inserir_linhas_info(doc, info_socicam, constatacao, analise_arpe):
    """Insere os blocos de Informação, Constatação e Análise de forma compacta (Seção 4)."""
    
    adicionar_paragrafo_info_compacta(doc, "Informação da SOCICAM: ", info_socicam)
    adicionar_paragrafo_info_compacta(doc, "Constatação: ", constatacao)
    adicionar_paragrafo_info_compacta(doc, "Análise da ARPE: ", analise_arpe)


def gerar_secao_nao_conformidades_constatadas(
    doc, row, nao_conformidades_df, fotos_dir, observacoes_df, recomendacoes_df
):
    """
    Gera a seção '4. RESULTADO DAS VISTORIAS DAS NÃO CONFORMIDADES CONSTATADAS'.
    """

    id_fisc = row["ID da Fiscalização"]
    
    nc_fisc = nao_conformidades_df[
        nao_conformidades_df["ID da Fiscalização"] == id_fisc
    ].copy() 
    
    quantidade_nc = len(nc_fisc)
    
    num_relatorio_original = row.get("CTR Original", "XX/XXXX") 
    num_monitoramento = row.get("Num Monitoramento", "Xº") 
    situacao_nc = row.get("Situacao NC", "PENDENTES")
    mes_ano = row.get("Mês/Ano Monitoramento", "MÊS/ANO") 

    adicionar_titulo_secao(doc, "4. RESULTADO DAS VISTORIAS DAS NÃO CONFORMIDADES CONSTATADAS")

    adicionar_paragrafo_justificado(
        doc,
        (
            f"Estão registrados, nos subitens a seguir, os resultados da verificação das ações desenvolvidas pela "
            f"SOCICAM para solucionar as Não Conformidades ainda pendentes, conforme o Quadro Resumo da "
            f"Situação das Não Conformidades Monitoradas ({mes_ano}) apresentado no {num_monitoramento} Relatório de "
            f"Monitoramento do Processo Arpe/CTR {num_relatorio_original} (Item 3). Ressalta-se que as {quantidade_nc} Não Conformidades "
            f"apontadas no referido Relatório de Monitoramento foram registradas como {situacao_nc}."
        )
    )

    if "Terminal" not in nc_fisc.columns:
        adicionar_paragrafo_justificado(
            doc, "⚠️ Coluna 'Terminal' não encontrada na planilha de não conformidades."
        )
        return

    num_terminal = 1
    
    for terminal, grupo_terminal in nc_fisc.groupby("Terminal"):
        
        grupo_terminal = grupo_terminal.sort_values(by="Nº")
        
        for numero_nc, grupo_nc in grupo_terminal.groupby("Nº"):
            
            nc_data = grupo_nc.iloc[0] 

            descricao = str(nc_data.get("Não Conformidade", "Descrição não disponível"))
            
            nc_titulo_identificador = str(nc_data.get("ID da não conformidade", str(numero_nc))) 

            info_socicam = str(nc_data.get("Informação SOCICAM", "Texto não disponível")).strip()
            constatacao = str(nc_data.get("Constatação", "Texto não disponível")).strip()
            analise_arpe = str(nc_data.get("Análise da Arpe", "Texto não disponível")).strip()
            
            _inserir_texto_nc(doc, num_terminal, terminal, nc_titulo_identificador, descricao)

            _inserir_linhas_info(doc, info_socicam, constatacao, analise_arpe)
            
            doc.add_paragraph() 
        
        num_terminal += 1
        
    return