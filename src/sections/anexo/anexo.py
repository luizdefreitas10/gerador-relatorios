import os 
from docx.document import Document
from pandas.core.series import Series as Row
from pandas import DataFrame
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Tuple

from utils import (
    adicionar_titulo_secao, 
    adicionar_duas_imagens_lado_a_lado, 
    aplicar_estilo_corpo,
    aplicar_estilo_titulo,
    adicionar_texto_contexto
) 


def _limpar_id_para_busca(nc_id: str) -> str:
    """Limpa o ID da NC para usá-lo como prefixo de arquivo (ex: '3.4 - CTR 02/2024' -> '3_4_CTR_02_2024')."""
    if not isinstance(nc_id, str):
        return ""
    
    nc_id = nc_id.replace(' ', '_').replace('-', '_').replace('.', '_').replace('/', '_')
    
    while '__' in nc_id:
        nc_id = nc_id.replace('__', '_')
    
    return nc_id.strip('_').upper()


def gerar_secao_anexo_fotos(doc: Document, row: Row, nao_conformidades_df: DataFrame, fotos_dir: str):
    """
    Gera a seção 'ANEXO - MEMORIAL FOTOGRÁFICO', buscando fotos no disco
    com base no 'ID da não conformidade' da planilha.
    """

    doc.add_page_break()

    caminho_fotos_monitoramento = fotos_dir 

    adicionar_titulo_secao(
        doc, 
        "ANEXO - MEMORIAL FOTOGRÁFICO - VISTORIAS REALIZADAS DE XX a XX/XX/XXXX",
        aplicar_sombra=True
    )
    id_fisc = row["ID da Fiscalização"]
    
    nc_fisc: DataFrame = nao_conformidades_df[
        nao_conformidades_df["ID da Fiscalização"] == id_fisc
    ].copy()
    
    if nc_fisc.empty:
        doc.add_page_break()
        doc.add_paragraph("Nenhuma não conformidade monitorada disponível.")
        return
 
    paragrafo_anexo = doc.add_paragraph()
    paragrafo_anexo.paragraph_format.space_after = Pt(12)
    run_paragrafo = paragrafo_anexo.add_run(
        "Observa-se o acompanhamento das Não Conformidades de acordo com os subitens do Relatório de "
        "Fiscalização Técnico-Operacional Arpe/CTR nº XX/XXXX."
    )
 
    aplicar_estilo_corpo(run_paragrafo, negrito=False)
    paragrafo_anexo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    
    texto_negrito = "NÃO CONFORMIDADES DOS TERMINAIS RODOVIÁRIOS DE PASSAGEIROS (TRP)"

    par_negrito = doc.add_paragraph()
    par_negrito.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par_negrito.paragraph_format.space_before = Pt(12) 
    par_negrito.paragraph_format.space_after = Pt(12) 

    run_negrito = par_negrito.add_run(texto_negrito)
    
    aplicar_estilo_titulo(run_negrito) 
    run_negrito.font.all_caps = True 
    
    
    #Pré-lista todos os arquivos da pasta para otimização
    try:
        arquivos_na_pasta = os.listdir(caminho_fotos_monitoramento)
    except FileNotFoundError:
        doc.add_paragraph(f"⚠️ ERRO: Pasta de fotos não encontrada em: {caminho_fotos_monitoramento}")
        return


    #Lógica de Geração de Fotos e Contexto 
    contexto_anterior: Tuple[str, str] = None 
    
    for terminal, grupo_terminal in nc_fisc.groupby("Terminal"):
        
        grupo_terminal = grupo_terminal.sort_values(by="Nº")
        
        for numero_nc, grupo_nc in grupo_terminal.groupby("Nº"):
            
            nc_data = grupo_nc.iloc[0] 
           
            terminal = str(nc_data.get("Terminal", "")).strip().upper() 
            nc_id_bruto = str(nc_data.get("ID da não conformidade", str(numero_nc))).strip() 
            constatacao = str(nc_data.get("Não Conformidade", "")).strip() 
            legenda_bruta = str(nc_data.get("Legenda da Foto", "")).strip() 
            
            prefixo_busca = _limpar_id_para_busca(nc_id_bruto)

            if not prefixo_busca:
                continue

            fotos_encontradas = [
                f for f in arquivos_na_pasta 
                if f.upper().startswith(f"{prefixo_busca}_") and f.lower().endswith(('.jpg', '.jpeg', '.png'))
            ]
            
            fotos_encontradas.sort() 
            
            legendas_lista = [leg.strip() for leg in legenda_bruta.split(';') if leg.strip()]
            
            contexto_atual_comparavel = (terminal, nc_id_bruto) 
            
            contexto_a_imprimir_para_fotos = None
            
            if contexto_atual_comparavel != contexto_anterior:

                contexto_para_texto = (
                    terminal, 
                    f" - Não Conformidade {nc_id_bruto}: {constatacao}"
                )
                
                if not fotos_encontradas:
  
                    adicionar_texto_contexto(doc, contexto_para_texto)
                
                else:
                    contexto_a_imprimir_para_fotos = (
                        terminal, 
                        nc_id_bruto,
                        constatacao
                    )

            if fotos_encontradas:
        
                fotos_em_pares = [
                    (fotos_encontradas[i], fotos_encontradas[i+1] if i + 1 < len(fotos_encontradas) else None) 
                    for i in range(0, len(fotos_encontradas), 2)
                ]
                
                for i, (foto1_nome, foto2_nome) in enumerate(fotos_em_pares):
                    
                    idx_foto1 = i * 2 
                    idx_foto2 = i * 2 + 1
                    
                    # Atribui a legenda se o índice existir na lista. Se não, usa uma string vazia ("").
                    legenda1 = legendas_lista[idx_foto1] if idx_foto1 < len(legendas_lista) else ""
                    # A segunda legenda só é verificada se houver uma segunda foto (foto2_nome is not None)
                    legenda2 = legendas_lista[idx_foto2] if foto2_nome and idx_foto2 < len(legendas_lista) else ""
     
                    adicionar_duas_imagens_lado_a_lado(
                        doc, 
                        caminho_fotos_monitoramento, 
                        foto1_nome,                   
                        legenda1, 
                        foto2_nome,                   
                        legenda2,
                        contexto_a_imprimir_para_fotos 
                    )
                   
                    contexto_a_imprimir_para_fotos = None 

            contexto_anterior = contexto_atual_comparavel